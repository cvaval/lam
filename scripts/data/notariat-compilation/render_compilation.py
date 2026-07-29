#!/usr/bin/env python3
"""
Rend le SOMMAIRE ANALYTIQUE et l'INDEX ALPHABÉTIQUE des 7 textes de la compilation
Notariat → deux .docx, une section par texte, au format des documents de la cliente.

Normalisations de l'index (l'IA sur-capitalise et varie la casse) :
  · fusion des variantes de CASSE d'un même libellé (« Acte Notarié » / « Acte notarié ») —
    sans ambiguïté possible : même chaîne une fois les accents et la casse neutralisés ;
  · capitalisation à la française — majuscule au premier mot seul, sauf institutions.
Les couples SINGULIER/PLURIEL et les sujets emboîtés sont SIGNALÉS, non fusionnés : les
rapprocher automatiquement confondrait des notions distinctes.

    python3 scripts/data/notariat-compilation/render_compilation.py
"""
import json
import os
import re
import unicodedata
from collections import defaultdict

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Pt, Inches

DIR = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.expanduser('~/Downloads')

INSTITUTIONS = {  # libellés déjà en apostrophe typographique
    'Ministère Public', 'Tribunal Civil', 'Tribunal de Première Instance', 'Juge de Paix',
    'Commissaire du Gouvernement', 'Conseil National de Gouvernement', 'Banque Nationale',
    'Conseil Supérieur du Notariat', 'Doyen du Tribunal Civil', 'Bâtonnier',
    "Secrétaire d'État de la Justice", 'Secrétaire d’État de la Justice',
    'Président de la République', 'Conseil d’État',
}

ORDRE = [
    'loi-1862-notariat', 'loi-1877-modificative', 'loi-1919-notariat', 'arrete-1919-examen',
    'decret-loi-1941-etude-vacante', 'decret-1974-nombre-notaires', 'decret-1986-nombre-notaires',
]


def apostrophe(s):
    """Apostrophe typographique unique — l'IA alterne « ' » et « ’ », ce qui dédoublait
    « Secrétaire d'État de la Justice »."""
    return s.replace("'", '\u2019')


def fold(s):
    return ''.join(c for c in unicodedata.normalize('NFD', apostrophe(s).lower())
                   if unicodedata.category(c) != 'Mn')


def cle_tri(s):
    f = fold(s)
    f = re.sub(r"^(l['’]|le |la |les |d['’]|de |du |des )", '', f)
    return re.sub(r'[^a-z0-9 ]', '', f)


def normaliser_casse(s):
    if s in INSTITUTIONS:
        return s
    mots = s.split(' ')
    out = [mots[0]]
    for m in mots[1:]:
        out.append(m if (m.isupper() and len(m) > 1) else (m[:1].lower() + m[1:] if m[:1].isupper() else m))
    return ' '.join(out)


def index_du_texte(donnees):
    """Construit l'index d'un texte : {sujet: {articles}} + journal des fusions de casse."""
    brut = defaultdict(set)
    for num, v in donnees.items():
        for s in v['s']:
            brut[apostrophe(s)].add(int(num))
    # fusion des variantes de casse/accent
    groupes = defaultdict(list)
    for s in brut:
        groupes[fold(s)].append(s)
    idx, journal = defaultdict(set), []
    for k, variantes in groupes.items():
        # on retient la forme la MOINS capitalisée à l'intérieur (usage français)
        retenue = sorted(variantes, key=lambda v: (sum(c.isupper() for c in v[1:]), v))[0]
        if len(variantes) > 1:
            journal.append(f'{" / ".join(sorted(variantes))} → {retenue}')
        for v in variantes:
            idx[normaliser_casse(retenue)] |= brut[v]
    return dict(idx), journal


def entete(doc, titre_gen, sous_titre, nota):
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)
    doc.add_paragraph().add_run(titre_gen).bold = True
    doc.add_paragraph().add_run(sous_titre).italic = True
    p = doc.add_paragraph()
    r = p.add_run(nota)
    r.font.size, r.italic = Pt(9), True


def main():
    textes = json.load(open(f'{DIR}/textes.json'))
    si = json.load(open(f'{DIR}/sommaire-index.json'))

    # ── SOMMAIRE ──
    ds = Document()
    entete(ds, 'SOMMAIRE ANALYTIQUE', 'Compilation Notariat — textes de 1862 à 1986',
           'Appareil éditorial — Les rubriques ci-dessous ne figurent pas au Journal Officiel ; '
           'elles ont été établies pour la lecture. Le décret-loi du 27 novembre 1969, présent dans '
           'la compilation, fait l’objet d’un document distinct.')
    for slug in ORDRE:
        t, d = textes[slug], si.get(slug, {})
        p = ds.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.add_run(t['titre']).bold = True
        for a in t['articles']:
            v = d.get(a['num'])
            if not v:
                continue
            lib = 'Article 1er' if a['num'] == '1' else f"Article {a['num']}"
            pe = ds.add_paragraph()
            pf = pe.paragraph_format
            pf.space_after, pf.left_indent, pf.first_line_indent = Pt(0), Inches(0.3), Inches(-0.3)
            pe.add_run(f'{lib}.- ').bold = True
            pe.add_run(v['r'])
    cs = os.path.join(OUT, 'Sommaire_Notariat_Compilation.docx')
    ds.save(cs)

    # ── INDEX ──
    di = Document()
    entete(di, 'INDEX ALPHABÉTIQUE', 'Compilation Notariat — textes de 1862 à 1986',
           'Nota. — Un index par texte ; les renvois sont faits aux articles du texte concerné. '
           'Ces entrées ne figurent pas au Journal Officiel : elles constituent un appareil éditorial.')
    total_sujets = total_refs = 0
    rapport = []
    final = {}
    for slug in ORDRE:
        t = textes[slug]
        idx, journal = index_du_texte(si.get(slug, {}))
        final[slug] = {k: sorted(v) for k, v in idx.items()}
        total_sujets += len(idx)
        total_refs += sum(len(v) for v in idx.values())
        p = di.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.add_run(t['titre']).bold = True
        lettre = None
        for s in sorted(idx, key=cle_tri):
            ini = (cle_tri(s)[:1] or '#').upper()
            if ini != lettre:
                lettre = ini
                pl = di.add_paragraph()
                pl.paragraph_format.space_before = Pt(6)
                pl.add_run(ini).bold = True
            refs = ' ; '.join(f'art. {x}' for x in sorted(idx[s]))
            pe = di.add_paragraph()
            pf = pe.paragraph_format
            pf.space_after, pf.left_indent, pf.first_line_indent = Pt(0), Inches(0.25), Inches(-0.25)
            pf.tab_stops.add_tab_stop(Inches(4.6), WD_TAB_ALIGNMENT.RIGHT)
            pe.add_run(f'{s}\t{refs}')
        couv = {n for r in idx.values() for n in r}
        rapport.append((slug, len(idx), sum(len(v) for v in idx.values()), len(couv), len(t['articles']), journal))
    ci = os.path.join(OUT, 'Index_Notariat_Compilation.docx')
    di.save(ci)
    # Index corrigé réinjectable par l'import (cf. commentaire du script du décret de 1969).
    json.dump(final, open(f'{DIR}/index-final.json', 'w'), ensure_ascii=False, indent=1)

    print(f'{"texte":32} sujets renvois couverture')
    for slug, ns, nr, couv, tot, journal in rapport:
        print(f'  {slug:30} {ns:5} {nr:7} {couv:6}/{tot}')
        for j in journal:
            print(f'{"":34} fusionné : {j}')
    print(f'\n  TOTAL : {total_sujets} sujets · {total_refs} renvois')
    print(f'  → {cs}')
    print(f'  → {ci}')


if __name__ == '__main__':
    main()
