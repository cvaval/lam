#!/usr/bin/env python3
"""
Met en forme le SOMMAIRE ANALYTIQUE et l'INDEX ALPHABÉTIQUE du décret-loi du 27 novembre
1969 sur le Notariat → deux .docx au format des documents de la cliente.

Normalisations appliquées à l'index (l'IA sur-capitalise et varie la casse) :
  · casse : seul le premier mot porte la majuscule — SAUF les institutions et noms propres
    (Tribunal Civil, Ministère Public, Commissaire du Gouvernement…), listés explicitement ;
  · fusions nominatives des doublons réels (« Acte Notarié »/« Acte notarié »,
    « Honoraire »/« Honoraires »…) — table explicite, jamais d'heuristique par ressemblance.

    python3 scripts/data/notariat-1969/render_sommaire_index.py
"""
import json
import os
import re
import unicodedata
from collections import defaultdict

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Pt, Inches

DIR = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.expanduser('~/Downloads')

TITRE = 'Décret-loi du 27 novembre 1969 harmonisant les dispositions de la Loi du 24 février 1919 sur le Notariat'
MONITEUR = 'Le Moniteur — Décret du 27 novembre 1969 (Dr. François Duvalier)'

# Sujets dont la capitalisation interne est VOULUE (institutions, noms propres).
INSTITUTIONS = {
    'Banque Nationale', 'Commissaire du Gouvernement', 'Conseil Supérieur du Notariat',
    'Doyen du Tribunal Civil', 'Juge de Paix', 'Ministère Public', 'Tribunal Civil',
    "Secrétaire d'État de la Justice", 'Secrétaire d’État de la Justice',
}

# Fusions nominatives : ancien libellé → libellé retenu (les renvois sont réunis).
FUSIONS = {
    'Acte Notarié': 'Acte notarié',
    'Dommages-Intérêts': 'Dommages-intérêts',
    'Dommage-intérêt': 'Dommages-intérêts',
    'Timbre Mobile Spécial': 'Timbre mobile spécial',
    'Honoraire': 'Honoraires',
    'Prestation de Serment': 'Serment',
    'Étude': 'Étude notariale',
    # « Ministère » (art. 2) désigne le MINISTÈRE DU NOTAIRE (« prêter leur ministère »),
    # non le Ministère Public : libellé désambiguïsé.
    'Ministère': 'Ministère du notaire',
}


def normaliser_casse(s):
    """Majuscule au premier mot seulement — sauf institutions et sigles."""
    if s in INSTITUTIONS:
        return s
    mots = s.split(' ')
    out = [mots[0]]
    for m in mots[1:]:
        # sigles et nombres romains conservés ; « d'Examen » → « d'examen »
        if m.isupper() and len(m) > 1:
            out.append(m)
        else:
            out.append(re.sub(r"([A-ZÀ-ÖØ-Þ])", lambda x: x.group(1).lower(), m, count=1)
                       if re.match(r"^[a-zà-öø-ÿ]*['’]?[A-ZÀ-ÖØ-Þ]", m) or m[:1].isupper() else m)
    return ' '.join(out)


def cle_tri(s):
    f = ''.join(c for c in unicodedata.normalize('NFD', s.lower()) if unicodedata.category(c) != 'Mn')
    return re.sub(r"[^a-z0-9 ]", '', f)


def charger():
    d = json.load(open(f'{DIR}/sommaire-index.json'))
    struct = json.load(open(f'{DIR}/structure.json'))
    return d, struct


def construire_index(d):
    brut = defaultdict(set)
    for n, v in d.items():
        for s in v['s']:
            brut[s].add(int(n))
    # 1. fusions nominatives  2. normalisation de casse  3. regroupement final
    fusionne = defaultdict(set)
    journal = []
    for s, refs in brut.items():
        cible = FUSIONS.get(s, s)
        if cible != s:
            journal.append(f'{s} → {cible}')
        fusionne[normaliser_casse(cible)] |= refs
    return dict(fusionne), journal


def rendre_index(idx, journal):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)
    doc.add_paragraph().add_run('INDEX ALPHABÉTIQUE').bold = True
    doc.add_paragraph().add_run(TITRE).italic = True
    doc.add_paragraph(f'({MONITEUR})')
    n = doc.add_paragraph()
    r = n.add_run('Nota. — Les renvois sont faits aux articles du Décret. Les entrées de cet index ne '
                  'figurent pas au Journal Officiel : elles constituent un appareil éditorial.')
    r.font.size, r.italic = Pt(9), True

    lettre = None
    for s in sorted(idx, key=cle_tri):
        ini = (cle_tri(s)[:1] or '#').upper()
        if ini != lettre:
            lettre = ini
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.add_run(ini).bold = True
        refs = ' ; '.join(f'art. {x}' for x in sorted(idx[s]))
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after, pf.left_indent, pf.first_line_indent = Pt(0), Inches(0.25), Inches(-0.25)
        pf.tab_stops.add_tab_stop(Inches(4.6), WD_TAB_ALIGNMENT.RIGHT)
        p.add_run(f'{s}\t{refs}')
    chemin = os.path.join(OUT, 'Index_Decret_Loi_Notariat_1969.docx')
    doc.save(chemin)
    return chemin


def rendre_sommaire(d, struct):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)
    doc.add_paragraph().add_run('SOMMAIRE ANALYTIQUE').bold = True
    doc.add_paragraph().add_run(TITRE).italic = True
    doc.add_paragraph(f'({MONITEUR})')
    n = doc.add_paragraph()
    r = n.add_run('Appareil éditorial — Les rubriques ci-dessous ne figurent pas au Journal Officiel ; '
                  'elles ont été établies pour la lecture. Le décret compte 80 articles.')
    r.font.size, r.italic = Pt(9), True

    # Position de chaque en-tête : rattaché à l'article qui l'ouvre (structure.json + corps).
    body = open(f'{DIR}/bodyOriginal.txt').read().split('\n')
    entetes = {t['label']: t for t in struct['toc']}
    courant = None
    for ligne in body:
        l = ligne.strip()
        if l in entetes:
            courant = entetes[l]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10 if courant['level'] == 1 else 6)
            run = p.add_run(l)
            run.bold = True
            if courant['level'] == 2:
                run.font.size = Pt(10)
            continue
        m = re.match(r'^Article\s+(\d{1,3})\.\s*—', l)
        if not m:
            continue
        num = m.group(1)
        v = d.get(num)
        if not v:
            continue
        libelle = 'Article 1er' if num == '1' else f'Article {num}'
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after, pf.left_indent, pf.first_line_indent = Pt(0), Inches(0.3), Inches(-0.3)
        p.add_run(f'{libelle}.- ').bold = True
        p.add_run(v['r'])
    chemin = os.path.join(OUT, 'Sommaire_Decret_Loi_Notariat_1969.docx')
    doc.save(chemin)
    return chemin


def main():
    d, struct = charger()
    idx, journal = construire_index(d)
    # L'index CORRIGÉ (fusions + capitalisation) est réécrit en JSON : sans cela l'import
    # reprenait la sortie BRUTE de l'IA et la plateforme affichait un index moins bon que
    # le .docx livré à la cliente.
    json.dump({k: sorted(v) for k, v in idx.items()}, open(f'{DIR}/index-final.json', 'w'),
              ensure_ascii=False, indent=1)
    ci = rendre_index(idx, journal)
    cs = rendre_sommaire(d, struct)
    nrefs = sum(len(v) for v in idx.values())
    couv = {n for r in idx.values() for n in r}
    print(f'index    : {len(idx)} sujets · {nrefs} renvois · couverture {len(couv)}/80 articles')
    for j in journal:
        print(f'           fusionné : {j}')
    print(f'           → {ci}')
    print(f'sommaire : {len(d)} rubriques · {len(struct["toc"])} en-têtes')
    print(f'           → {cs}')


if __name__ == '__main__':
    main()
