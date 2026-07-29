#!/usr/bin/env python3
"""
Consolide puis met en forme les index alphabétiques générés → fichiers .docx au format
exact de ceux de la cliente (« INDEX ALPHABÉTIQUE », note de renvoi, en-têtes de lettre,
sujet + TABULATION + renvois).

Consolidation : un sujet « satellite » (dont le libellé contient un autre sujet et dont les
renvois sont TOUS déjà couverts par lui) n'apporte aucun point d'entrée nouveau — il est
fusionné dans le sujet principal. « Certificat électronique qualifié » qui renvoie à un
AUTRE article est en revanche conservé : c'est une entrée utile.

    python3 scripts/data/electronique-2015-2025/render_index_docx.py
"""
import json
import os
import re
import unicodedata

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Pt, Inches

DIR = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.expanduser('~/Downloads')

TEXTES = {
    'decret-2015-signature': (
        'Décret du 9 décembre 2015 sur la signature électronique',
        'Le Moniteur, 171e Année, No. 20, vendredi 29 janvier 2016',
        'Les renvois sont faits aux articles du Décret. La mention « C. civ. » indique '
        'l’article du Code Civil dans sa rédaction issue du Décret.',
        'Index_Decret_Signature_Electronique_2015.docx',
    ),
    'decret-2016-administration': (
        'Décret du 6 janvier 2016 reconnaissant le droit de tout administré à s’adresser à '
        'l’administration publique par des moyens électroniques',
        'Le Moniteur, 171e Année, No. 20, vendredi 29 janvier 2016',
        'Les renvois sont faits aux articles du Décret.',
        'Index_Decret_Administration_Electronique_2016.docx',
    ),
    'loi-2017-echanges': (
        'Loi du 14 février 2017 sur les échanges électroniques',
        'Le Moniteur, 172e Année, Spécial No 12, 11 avril 2017',
        'Les renvois sont faits aux articles de la Loi.',
        'Index_Loi_2017_Echanges_Electroniques.docx',
    ),
}


# Corrections ÉDITORIALES explicites (relevées à la relecture de la sortie IA). Table
# nominative plutôt qu'heuristique : une fusion automatique par ressemblance rapprocherait
# des notions distinctes (« Certificat électronique » ≠ « Certificat électronique qualifié »).
#   renommer : ancien libellé → nouveau (les renvois sont fusionnés si la cible existe déjà)
#   supprimer : sujets sans valeur d'index
CORRECTIONS = {
    'decret-2015-signature': {
        # Deux libellés pour UNE seule notion — l'index en devenait incohérent.
        'renommer': {
            'Prestataire de services de certification':
                'Prestataire de services de certification électronique',
            'Responsabilité': 'Responsabilité du notaire',  # art. 5 : responsabilité des officiers publics
        },
        'supprimer': [],
    },
    'decret-2016-administration': {'renommer': {}, 'supprimer': []},
    'loi-2017-echanges': {
        'renommer': {
            # Anglicisme « (Definition) » produit par le modèle : la définition d'une notion
            # se range sous la notion elle-même, pas sous une entrée séparée.
            'Message de données (Definition)': 'Message de données',
            'Expéditeur (Definition)': 'Expéditeur',
            'Échange de données informatisées (Definition)':
                'Échange de données informatisées (EDI)',
            'Relation extra contractuelle': 'Relation extracontractuelle',
        },
        # art. 3 déjà couvert par « Interprétation de la loi » et « Uniformité d'application ».
        'supprimer': ['Principe général'],
    },
}


def corriger(slug, idx):
    """Applique la table de corrections. Renvoie (index, journal)."""
    conf = CORRECTIONS.get(slug, {})
    journal = []
    for ancien, nouveau in (conf.get('renommer') or {}).items():
        if ancien not in idx:
            continue
        refs = set(idx.pop(ancien))
        fusion = nouveau in idx
        idx[nouveau] = sorted(set(idx.get(nouveau, [])) | refs, key=cle_ref)
        journal.append(f'{"fusionné" if fusion else "renommé"} : « {ancien} » → « {nouveau} »')
    for s in conf.get('supprimer') or []:
        if s in idx:
            idx.pop(s)
            journal.append(f'supprimé : « {s} »')
    return idx, journal


def fold(s):
    return ''.join(c for c in unicodedata.normalize('NFD', s.lower())
                   if unicodedata.category(c) != 'Mn')


def cle_tri(s):
    """Tri français : accents ignorés, apostrophes et articles élidés neutralisés."""
    f = fold(s)
    f = re.sub(r"^(l['’]|le |la |les |d['’]|de |du |des )", '', f)
    return re.sub(r"[^a-z0-9 ]", '', f)


def cle_ref(r):
    return [int(p) for p in r.split('-')]


def consolider(idx):
    """Fusionne les satellites sans apport. Renvoie (index consolidé, journal des fusions)."""
    sujets = sorted(idx, key=lambda s: len(fold(s)))  # les plus courts d'abord = candidats parents
    supprimes, journal = set(), []
    for court in sujets:
        fc = fold(court)
        for long in sujets:
            if long == court or long in supprimes or court in supprimes:
                continue
            fl = fold(long)
            if fc == fl or fc not in fl:
                continue
            # frontière de mot : « acte » ne doit pas « absorber » « caractère »
            if not re.search(rf'(^|\W){re.escape(fc)}(\W|$)', fl):
                continue
            if set(idx[long]) <= set(idx[court]):
                journal.append((long, court))
                supprimes.add(long)
    out = {s: r for s, r in idx.items() if s not in supprimes}
    return out, journal


def rendre(slug, idx, titre, moniteur, note, nom_fichier):
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(10.5)

    h = doc.add_paragraph()
    h.add_run('INDEX ALPHABÉTIQUE').bold = True
    p = doc.add_paragraph()
    p.add_run(titre).italic = True
    doc.add_paragraph(f'({moniteur})')
    n = doc.add_paragraph()
    r = n.add_run('Nota. — ' + note + ' Les entrées de cet index ne figurent pas au Journal '
                  'Officiel : elles constituent un appareil éditorial.')
    r.font.size = Pt(9)
    r.italic = True

    lettre_courante = None
    for sujet in sorted(idx, key=cle_tri):
        initiale = (cle_tri(sujet)[:1] or '#').upper()
        if initiale != lettre_courante:
            lettre_courante = initiale
            pl = doc.add_paragraph()
            pl.paragraph_format.space_before = Pt(8)
            pl.add_run(initiale).bold = True
        refs = ' ; '.join('art. ' + x for x in sorted(idx[sujet], key=cle_ref))
        pe = doc.add_paragraph()
        pf = pe.paragraph_format
        pf.space_after = Pt(0)
        pf.left_indent = Inches(0.25)
        pf.first_line_indent = Inches(-0.25)
        pf.tab_stops.add_tab_stop(Inches(4.6), WD_TAB_ALIGNMENT.RIGHT)
        pe.add_run(f'{sujet}\t{refs}')

    chemin = os.path.join(OUT, nom_fichier)
    doc.save(chemin)
    return chemin


def main():
    for slug, (titre, moniteur, note, nom) in TEXTES.items():
        f = f'{DIR}/index-{slug}.json'
        idx = json.load(open(f))
        avant = len(idx)
        idx, corr = corriger(slug, idx)
        idx, journal = consolider(idx)
        json.dump(idx, open(f, 'w'), ensure_ascii=False, indent=1)
        chemin = rendre(slug, idx, titre, moniteur, note, nom)
        nrefs = sum(len(v) for v in idx.values())
        print(f'{slug:30} {avant:3} → {len(idx):3} sujets ({avant - len(idx)} traités) · {nrefs} renvois')
        for c in corr:
            print(f'{"":32} {c}')
        for long, court in journal[:3]:
            print(f'{"":32} absorbé : « {long} » → « {court} »')
        print(f'{"":32} → {chemin}')


if __name__ == '__main__':
    main()
